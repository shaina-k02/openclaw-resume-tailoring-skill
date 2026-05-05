#!/usr/bin/env python3
"""Generate sample resume files (PDF and DOCX) for eval test cases.

Each file is a different person to make evals more realistic:
- sample_resume.pdf  → Alex Johnson (backend/infra engineer) - used for eval 1 (Stripe)
- sample_resume.docx → Jordan Rivera (ML/AI engineer) - used for eval 2 (ML startup)
- sample_resume.tex  → Maya Chen (full-stack engineer) - already exists, used for eval 3 (Microsoft)
"""

import os


# ── Alex Johnson: Backend / Infrastructure (PDF - eval 1: Stripe Data Engineer) ──

ALEX_RESUME = {
    "name": "Alex Johnson",
    "email": "alex.johnson@email.com",
    "phone": "(555) 123-4567",
    "linkedin": "linkedin.com/in/alexjohnson",
    "github": "github.com/alexj",
    "location": "San Francisco, CA",
    "summary": (
        "Software engineer with 6 years of experience building scalable backend systems, "
        "data pipelines, and cloud-native applications. Strong background in Python, Java, "
        "and Go with deep expertise in distributed systems, cloud infrastructure, and DevOps. "
        "Passionate about clean architecture, performance optimization, and mentoring junior engineers."
    ),
    "experience": [
        {
            "title": "Senior Software Engineer",
            "company": "CloudScale Technologies",
            "dates": "Mar 2024 - Present",
            "bullets": [
                "Led the design and implementation of a microservices migration, decomposing a monolithic Django app into 12 independently deployable services using Go and gRPC",
                "Architected an event-driven data pipeline processing 2M+ events/day using Apache Kafka, Apache Flink, and Amazon Kinesis",
                "Reduced infrastructure costs by 35% through autoscaling policies, spot instance strategies, and right-sizing EC2/EKS workloads",
                "Mentored a team of 4 junior engineers through weekly 1:1s, code reviews, and architecture design sessions",
                "Implemented distributed tracing with OpenTelemetry and Jaeger, reducing MTTR for production incidents by 60%",
            ],
        },
        {
            "title": "Software Engineer",
            "company": "TechCorp Inc.",
            "dates": "Jan 2022 - Feb 2024",
            "bullets": [
                "Designed and implemented RESTful APIs serving 50K+ daily active users using Python and FastAPI",
                "Built data processing pipelines using Apache Kafka and PostgreSQL, handling 500K+ records daily",
                "Reduced API response times by 40% through query optimization, Redis caching, and connection pooling",
                "Implemented CI/CD pipelines using GitHub Actions and Docker, reducing deployment time from 45 min to 8 min",
                "Developed an internal monitoring dashboard using Grafana and Prometheus across 15 services",
            ],
        },
        {
            "title": "Junior Developer",
            "company": "StartupXYZ",
            "dates": "Jun 2020 - Dec 2021",
            "bullets": [
                "Developed full-stack web applications using React, TypeScript, and Node.js serving 10K+ monthly users",
                "Built automated testing suites achieving 85% code coverage using Jest, Cypress, and pytest",
                "Integrated third-party APIs for payment processing (Stripe), notifications (Twilio), and analytics (Segment)",
                "Migrated legacy jQuery frontend to React, improving page load times by 55%",
            ],
        },
        {
            "title": "Software Engineering Intern",
            "company": "DataDriven Corp.",
            "dates": "May 2019 - Aug 2019",
            "bullets": [
                "Built an ETL pipeline using Python and Apache Airflow to ingest data from 5 external sources",
                "Created data visualizations using D3.js and Tableau for the executive analytics dashboard",
                "Wrote unit and integration tests for the data validation layer, catching 12 data quality issues",
            ],
        },
    ],
    "education": "B.S. Computer Science — State University (May 2020), GPA: 3.6/4.0, Dean's List",
    "skills": [
        "Python", "Java", "Go", "JavaScript", "TypeScript", "SQL",
        "FastAPI", "Django", "React", "Node.js", "Spring Boot",
        "PostgreSQL", "MongoDB", "Redis", "DynamoDB", "Elasticsearch",
        "AWS (EC2, S3, Lambda, EKS, RDS, Kinesis)", "Docker", "Kubernetes", "Terraform", "GitHub Actions",
        "Apache Kafka", "Apache Flink", "Apache Airflow", "Spark",
        "Prometheus", "Grafana", "OpenTelemetry", "Git", "Linux", "Agile/Scrum",
    ],
    "projects": [
        ("FastAPI Rate Limiter Middleware (Open Source)", "Production-grade rate-limiting middleware for FastAPI. 500+ GitHub stars, used by 30+ companies."),
        ("Distributed Task Queue (Personal)", "Lightweight task queue in Go with prioritization, retry logic, and dead letter queues. Uses Redis."),
        ("Real-Time Collaborative Editor (HackSF 2023 Winner)", "CRDT-based collaborative editor. Won 1st place out of 120 teams. React, Node.js, Y.js."),
    ],
    "certifications": [
        "AWS Certified Solutions Architect - Associate (2023)",
        "Certified Kubernetes Administrator (CKA) (2024)",
    ],
}


# ── Jordan Rivera: ML / AI Engineer (DOCX - eval 2: ML startup) ──

JORDAN_RESUME = {
    "name": "Jordan Rivera",
    "email": "jordan.rivera@email.com",
    "phone": "(312) 555-8294",
    "linkedin": "linkedin.com/in/jordanrivera",
    "github": "github.com/jrivera-ml",
    "location": "Chicago, IL",
    "summary": (
        "Machine learning engineer with 3 years of experience building and deploying production "
        "ML systems. Skilled in Python, PyTorch, and cloud ML infrastructure. Experience spans "
        "NLP, computer vision, and recommendation systems. Strong foundation in software engineering "
        "practices including CI/CD, containerization, and monitoring for ML pipelines."
    ),
    "experience": [
        {
            "title": "Machine Learning Engineer",
            "company": "Predictive AI Labs",
            "dates": "Sep 2023 - Present",
            "bullets": [
                "Designed and deployed a real-time recommendation engine serving 5M+ daily predictions using PyTorch, FastAPI, and Redis",
                "Built an automated ML pipeline with MLflow, Airflow, and AWS SageMaker for model training, evaluation, and deployment",
                "Fine-tuned LLaMA 2 and Mistral models for domain-specific text classification, achieving 92% F1 score on internal benchmarks",
                "Implemented A/B testing framework for ML models using feature flags and statistical significance testing",
                "Reduced model inference latency by 65% through ONNX runtime optimization and batched prediction serving",
                "Collaborated with product teams to define ML success metrics and translate business requirements into model objectives",
            ],
        },
        {
            "title": "Data Scientist",
            "company": "UrbanTech Solutions",
            "dates": "Jul 2021 - Aug 2023",
            "bullets": [
                "Built a computer vision pipeline using OpenCV and TensorFlow for automated defect detection on manufacturing lines, reducing manual inspection time by 40%",
                "Developed NLP models for customer support ticket classification using BERT and spaCy, routing 70% of tickets automatically",
                "Created interactive Streamlit dashboards for non-technical stakeholders to explore model predictions and data trends",
                "Managed data labeling workflows using Label Studio, coordinating a team of 6 annotators across 3 projects",
                "Published internal technical blog posts on ML best practices adopted by the broader engineering team",
            ],
        },
        {
            "title": "Research Assistant",
            "company": "Northwestern University NLP Lab",
            "dates": "Sep 2019 - Jun 2021",
            "bullets": [
                "Co-authored a paper on few-shot text classification published at EMNLP 2021",
                "Implemented and benchmarked transformer-based models (BERT, RoBERTa, GPT-2) for sentiment analysis tasks",
                "Built a data augmentation pipeline that increased training data diversity by 3x using back-translation and paraphrasing",
                "Presented research findings at 2 university symposiums and 1 industry workshop",
            ],
        },
    ],
    "education": "M.S. Computer Science (ML Specialization) — Northwestern University (Jun 2021), GPA: 3.8/4.0\nB.S. Mathematics & Computer Science — University of Illinois Chicago (May 2019), GPA: 3.5/4.0",
    "skills": [
        "Python", "SQL", "Bash", "R",
        "PyTorch", "TensorFlow", "scikit-learn", "Hugging Face Transformers", "spaCy", "OpenCV",
        "FastAPI", "Streamlit", "Flask",
        "PostgreSQL", "Redis", "Pinecone", "Elasticsearch",
        "AWS (SageMaker, EC2, S3, Lambda)", "Docker", "Kubernetes", "MLflow", "Airflow", "Weights & Biases",
        "Git", "Linux", "Jupyter", "pandas", "NumPy", "ONNX Runtime",
    ],
    "projects": [
        ("RAG-powered Document Q&A (Personal)", "Built a retrieval-augmented generation system using LangChain, Pinecone, and GPT-4 for querying technical documentation. Supports PDF ingestion, chunking, and semantic search."),
        ("Arxiv Paper Recommender (Open Source)", "Content-based recommendation engine for ML papers using sentence-transformers and FAISS. 200+ GitHub stars."),
        ("Real-Time Emotion Detection (Hackathon)", "Webcam-based emotion detection using a fine-tuned ResNet model and OpenCV. Won 2nd place at ChiHack 2022."),
    ],
    "certifications": [
        "AWS Certified Machine Learning - Specialty (2024)",
        "DeepLearning.AI TensorFlow Developer Certificate (2022)",
    ],
}


def generate_pdf(output_path):
    """Generate Alex Johnson's PDF resume."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    data = ALEX_RESUME
    doc = SimpleDocTemplate(output_path, pagesize=letter,
                            topMargin=0.5*inch, bottomMargin=0.5*inch,
                            leftMargin=0.75*inch, rightMargin=0.75*inch)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle("Title", parent=styles["Title"], fontSize=18, spaceAfter=4)
    contact_style = ParagraphStyle("Contact", parent=styles["Normal"], fontSize=10, alignment=1, spaceAfter=12)
    heading_style = ParagraphStyle("Head", parent=styles["Heading2"], fontSize=13, spaceAfter=6, spaceBefore=12)
    sub_style = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=11, spaceBefore=8, spaceAfter=2)
    body = styles["Normal"]
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=20, bulletIndent=10, spaceBefore=2)

    story = []
    story.append(Paragraph(data["name"], title_style))
    story.append(Paragraph(f"{data['email']} | {data['phone']} | {data['linkedin']} | {data['github']}<br/>{data['location']}", contact_style))

    story.append(Paragraph("Professional Summary", heading_style))
    story.append(Paragraph(data["summary"], body))

    story.append(Paragraph("Experience", heading_style))
    for job in data["experience"]:
        story.append(Paragraph(f"<b>{job['title']} — {job['company']}</b>  ({job['dates']})", sub_style))
        for b in job["bullets"]:
            story.append(Paragraph(f"• {b}", bullet))
        story.append(Spacer(1, 4))

    story.append(Paragraph("Education", heading_style))
    story.append(Paragraph(data["education"], body))

    story.append(Paragraph("Technical Skills", heading_style))
    story.append(Paragraph(", ".join(data["skills"]), body))

    story.append(Paragraph("Projects", heading_style))
    for name, desc in data["projects"]:
        story.append(Paragraph(f"<b>{name}</b>", sub_style))
        story.append(Paragraph(desc, body))

    story.append(Paragraph("Certifications", heading_style))
    for c in data["certifications"]:
        story.append(Paragraph(f"• {c}", bullet))

    doc.build(story)
    print(f"Created: {output_path}")


def generate_docx(output_path):
    """Generate Jordan Rivera's DOCX resume."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = JORDAN_RESUME
    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data["name"])
    run.bold = True
    run.font.size = Pt(18)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(f"{data['email']} | {data['phone']} | {data['linkedin']} | {data['github']}\n{data['location']}").font.size = Pt(10)

    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(data["summary"])

    doc.add_heading("Experience", level=2)
    for job in data["experience"]:
        p = doc.add_paragraph()
        p.add_run(f"{job['title']} — {job['company']}").bold = True
        p.add_run(f"  ({job['dates']})")
        for b in job["bullets"]:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Education", level=2)
    for line in data["education"].split("\n"):
        doc.add_paragraph(line)

    doc.add_heading("Technical Skills", level=2)
    doc.add_paragraph(", ".join(data["skills"]))

    doc.add_heading("Projects", level=2)
    for name, desc in data["projects"]:
        p = doc.add_paragraph()
        p.add_run(name).bold = True
        doc.add_paragraph(desc)

    doc.add_heading("Certifications", level=2)
    for c in data["certifications"]:
        doc.add_paragraph(c, style="List Bullet")

    doc.save(output_path)
    print(f"Created: {output_path}")


if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    files_dir = os.path.join(script_dir, "..", "evals", "files")
    os.makedirs(files_dir, exist_ok=True)

    try:
        import docx
    except ImportError:
        os.system("pip install python-docx --break-system-packages -q")
    try:
        import reportlab
    except ImportError:
        os.system("pip install reportlab --break-system-packages -q")

    generate_pdf(os.path.join(files_dir, "sample_resume.pdf"))
    generate_docx(os.path.join(files_dir, "sample_resume.docx"))
    print("\nSample eval files generated successfully.")
    print("  sample_resume.pdf  → Alex Johnson (backend/infra)")
    print("  sample_resume.docx → Jordan Rivera (ML/AI)")
    print("  sample_resume.tex  → Maya Chen (full-stack) [already exists]")
